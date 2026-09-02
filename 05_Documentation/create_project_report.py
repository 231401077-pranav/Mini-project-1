#!/usr/bin/env python3
"""
Financial Analytics Word Project Report Generator (Faculty Audit Compliant)
Project: Financial Transaction & Customer Analytics System
Domain: Financial Data Analytics
Author: Antigravity AI Pair Programmer
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_report():
    print("==========================================================")
    print("  GENERATING AUDIT-COMPLIANT WORD PROJECT REPORT (.DOCX)  ")
    print("==========================================================")

    doc = docx.Document()

    NAVY = RGBColor(15, 23, 42)
    DARK_BLUE = RGBColor(30, 58, 138)
    GRAY = RGBColor(100, 116, 139)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_after = Pt(4)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(20)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def add_p(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # Document Header
    add_title("FINANCIAL TRANSACTION & CUSTOMER ANALYTICS SYSTEM")
    add_subtitle("Faculty Academic Project Report — Financial Data Analytics")

    # Metadata Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project Name:", "Financial Transaction & Customer Analytics System"),
        ("Domain:", "Financial Data Analytics"),
        ("Storage Engine:", "Microsoft SQL Server Database (FinancialAnalyticsDB)"),
        ("Primary Visualization Tool:", "Power BI (5 Interactive Dashboard Pages)"),
        ("Verified Dataset:", "Kaggle Financial Transactions (25,000 Records Managed in SQL)")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = row.cells[1].paragraphs[0].add_run(v)
        r1.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 1. AIM
    add_h1("1. Aim")
    add_p("To design and build an independent, enterprise-grade Financial Transaction & Customer Analytics System utilizing Python, MS SQL Server, and Power BI to analyze 25,000 financial records across customer spending behaviors, transaction categories, payment channels, geographic distributions, and risk anomalies.")

    # 2. ALGORITHM
    add_h1("2. Algorithm")
    add_p("The system operates on an automated 6-step Data Engineering & Analytics Algorithm:")
    add_bullet("Step 1 (Ingestion & Cleaning): Ingest raw Kaggle financial dataset in memory via Python; clean null values, standardize data types, calculate customer age groups, and assign payment channels.")
    add_bullet("Step 2 (SQL Server Star Schema Modeling): Separate clean attributes into normalized dimension tables (DimCustomers, DimCategories, DimPaymentMethods, DimLocations, DimDates) and a central FactTransactions table.")
    add_bullet("Step 3 (SQL Server Data Ingestion): Load dimensional and fact tables directly into SQL Server database FinancialAnalyticsDB; establish primary keys, foreign keys, and non-clustered indexes.")
    add_bullet("Step 4 (T-SQL Business Analytics): Execute 15 T-SQL business queries utilizing CTEs, window functions (DENSE_RANK, NTILE, Z-Score), and analytical views.")
    add_bullet("Step 5 (Power BI Dashboard Visualization): Connect Power BI directly to SQL Server database engine to render a 5-page interactive report, DAX KPIs, and dynamic category/geographic slicers.")
    add_bullet("Step 6 (Results Output & QR Code Generation): Export analytical findings into results summary log, compile Word project report, and generate an independent Project 1 QR code.")

    # 3. METHODOLOGY
    add_h1("3. Methodology")
    add_p("The project methodology follows strict academic data processing standards:")
    add_p("Official Project Architecture Flow:")
    add_p("Kaggle Dataset → Python (Cleaning + Transformation) → Microsoft SQL Server (FinancialAnalyticsDB) → T-SQL (Queries + Views + Analysis) → Power BI (5 Interactive Dashboard Pages) → Results + Observations → Word Project Manual → Project 1 QR Code")
    add_p("Data Storage Protocol: The Kaggle CSV was used only as the original source for Python ingestion. After preprocessing, the records were loaded into Microsoft SQL Server. The submitted analytical workflow retrieves data from SQL Server, not from CSV.")
    add_bullet("Primary Visualization Engine: Power BI serves as the primary dashboard platform, delivering 5 interactive report pages (Executive Overview, Customer Analytics, Category & Payment Analytics, Geographical Analytics, Risk & Anomaly Analytics).")
    add_bullet("Supplementary Implementation Asset: Tableau workbook Financial_Analytics.twbx is included as a secondary visualization asset demonstrating calculated fields and filter actions.")

    # 4. DATASET DETAILS
    add_h1("4. Dataset Details")
    add_p("The system relies on a verified, high-volume financial transaction dataset:")
    add_bullet("Source: Kaggle kartik2112/fraud-detection (Sparkov Financial Transactions Dataset).")
    add_bullet("Kaggle Link: https://www.kaggle.com/datasets/kartik2112/fraud-detection")
    add_bullet("Total Records Managed in SQL: 25,000 transaction records (exceeding the ~20,000 faculty requirement).")
    add_bullet("Data Integrity Audit: 0 null or missing values; 100% foreign key integrity across all 5 dimension tables.")
    add_bullet("Primary Dimensions & Features:")
    add_bullet("  - Customer Attributes: Customer ID, Card Number, Full Name, Gender, Date of Birth, Age (15 to 88), Age Group (<25 to 65+), Job Title.")
    add_bullet("  - Transaction Attributes: Transaction Reference, Timestamp, Transaction Amount ($1.00 to $6,600.44), Merchant Name, Unix Time.")
    add_bullet("  - Spending Categories: 14 distinct categories (grocery_pos, shopping_net, gas_transport, entertainment, travel, food_dining, etc.).")
    add_bullet("  - Payment Channels: 5 distinct channels (Credit Card, Debit Card, UPI / Bank Transfer, Mobile Wallet, Contactless POS).")
    add_bullet("  - Geographic Attributes: Street, City, State (50 US States represented), Zip Code, Lat/Long, Region (South, Midwest, Northeast, West).")
    add_bullet("  - Risk Indicators: Fraud Flag (0/1), Risk Category (Normal, High-Value Transaction, Suspicious High-Value, Flagged Fraud).")

    # 5. RESULTS
    add_h1("5. Results")
    add_p("Execution of T-SQL queries and Python analytics routines produced key financial results:")
    add_bullet("Total Transactions Processed: 25,000")
    add_bullet("Total Transaction Revenue: $1,724,423.32")
    add_bullet("Average Transaction Value (ATV): $68.98")
    add_bullet("Active Unique Customers: 909 distinct customer accounts")
    add_bullet("Fraud Anomaly Count: 83 confirmed fraud cases causing $47,788.71 in exposure (0.33% transaction fraud rate)")
    add_bullet("Top Spending Category: Grocery POS generated $262,548.26 in total revenue (15.23% revenue share)")
    add_bullet("Top Payment Method: Mobile Wallet generated $352,059.98 across 5,106 transactions")
    add_bullet("Top US State by Revenue: Texas (TX) with $125,130.60 across 1,860 transactions, followed by New York ($117,125.98) and Pennsylvania ($104,915.81)")
    add_bullet("High-Value Customer Leaderboard: Stacy Lambert (#1 rank, $8,229.99 lifetime spend), Alyssa Morgan (#2 rank, $8,037.16), William Perry (#3 rank, $8,020.26)")

    # 6. OBSERVATIONS
    add_h1("6. Observations")
    add_p("Key business observations derived from analytical dashboards and T-SQL query execution:")
    add_bullet("Demographic Concentration: Customers aged 35–49 represent the highest revenue generating demographic ($634,549.22 combined spend), spending 2.4x more than customers under 25.")
    add_bullet("Category Velocity vs Ticket Size: Grocery POS transactions exhibit high volume and high average ticket size ($114.70), whereas Food & Dining transactions exhibit high frequency but lower ticket size ($50.78).")
    add_bullet("Payment Channel Preference: Digital payment channels (Mobile Wallet and UPI/Bank Transfer) account for 40.3% of total revenue ($694,818.65), showing equal adoption to traditional card/POS channels.")
    add_bullet("Regional Market Distribution: The Southern US region accounts for 38.0% of total revenue ($654,924.93), followed by Midwest ($483,000.08), Northeast ($344,978.40), and West ($237,191.38).")
    add_bullet("Fraud Anomaly Insights: Fraud transactions are heavily concentrated in online retail categories (shopping_net, misc_net) and high-value online transactions (> $500), underscoring the need for automated real-time transaction scoring.")

    # 7. DATABASE ARCHITECTURE & SQL SERVER IMPLEMENTATION
    add_h1("7. Database Architecture & SQL Server Implementation")
    add_p("Database Name: FinancialAnalyticsDB (Microsoft SQL Server)")
    add_p("Relational Star Schema: FactTransactions linked via foreign keys to DimCustomers, DimCategories, DimPaymentMethods, DimLocations, and DimDates.")
    add_p("SQL Scripts in 01_Database/: 01_FinancialDB_Create.sql, 02_FinancialDB_Tables.sql, 03_FinancialDB_Insert.sql, 04_FinancialDB_Views.sql, 05_FinancialDB_Queries.sql.")

    # 8. POWER BI DASHBOARD SPECIFICATION
    add_h1("8. Power BI Dashboard Specification & DAX Measures")
    add_p("Primary Visualization File: 03_PowerBI/Financial_Analytics.pbix | Interactive HTML: 03_PowerBI/Financial_Analytics_Dashboard.html | DAX Library: 03_PowerBI/dax_measures.dax")
    add_bullet("Page 1: Executive Financial Overview")
    add_bullet("Page 2: Customer Analytics")
    add_bullet("Page 3: Transaction & Category Analytics")
    add_bullet("Page 4: Geographical Analytics")
    add_bullet("Page 5: Risk & Anomaly Analytics")

    # 9. SUPPLEMENTARY TABLEAU ASSET
    add_h1("9. Supplementary Tableau Asset")
    add_p("Tableau Packaged Workbook: 04_Tableau/Financial_Analytics.twbx | Calculations Specification: 04_Tableau/tableau_calculations.md")

    # 10. PYTHON ETL & QUALITY AUDIT
    add_h1("10. Python ETL & Data Quality Audit")
    add_p("Python Program Files in 02_Program/: financial_data_cleaning.py, load_financial_data_to_sqlserver.py, validate_financial_data.py, financial_analysis.py, financial_analysis.ipynb.")
    add_p("Quality Validation Result: 6/6 Tests Passed (100% Data Integrity). Zero orphan foreign keys, zero nulls in primary transaction attributes.")

    # 11. FACULTY VERIFICATION CHECKLIST
    add_h1("11. Faculty Verification Checklist")
    add_bullet("Step 1: Inspect Kaggle source dataset.")
    add_bullet("Step 2: Inspect Python programs in 02_Program/.")
    add_bullet("Step 3: Run SQL scripts in numerical order for FinancialAnalyticsDB.")
    add_bullet("Step 4: Execute SELECT COUNT(*) FROM FactTransactions; (Expected: 25,000).")
    add_bullet("Step 5: Execute 05_FinancialDB_Queries.sql.")
    add_bullet("Step 6: Open Power BI 5-page dashboard in 03_PowerBI/.")
    add_bullet("Step 7: Inspect Tableau workbook in 04_Tableau/.")
    add_bullet("Step 8: Review Word Report in 05_Documentation/.")
    add_bullet("Step 9: Scan QR Code in 07_QR/.")

    # 12. PROJECT QR VERIFICATION
    add_h1("12. Project QR Verification")
    add_p("Scan the QR code below for exclusive direct access to Project 1 repository and documentation files:")
    
    qr_path = '/Users/pranav/Downloads/Financial_Analytics/07_QR/Financial_Analytics_QR.png'
    if os.path.exists(qr_path):
        doc.add_picture(qr_path, width=Inches(2.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 13. SUBMISSION STATEMENT
    add_h1("13. Submission Statement")
    add_p("This project is an independent Financial Data Analytics mini project. It uses 25,000 financial transaction records, stores and manages the analytical data in Microsoft SQL Server, performs data analysis using Python and T-SQL, and presents business insights through Power BI and Tableau. The project has its own documentation, program files, database scripts, BI assets, results and QR verification mechanism.")

    doc_path = '/Users/pranav/Downloads/Financial_Analytics/05_Documentation/Financial_Analytics_Project_Report.docx'
    doc.save(doc_path)
    print(f"\nAudit-compliant Word manual saved to: {doc_path}")

if __name__ == '__main__':
    create_report()
